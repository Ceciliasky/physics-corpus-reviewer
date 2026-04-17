import streamlit as st
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import time
from datetime import datetime
from pathlib import Path
from review_module import review_text, call_deepseek

# ================== 页面配置 ==================
st.set_page_config(page_title="物理教材智能助手", layout="wide")

# ================== 初始化 session_state ==================
if "history" not in st.session_state:
    st.session_state.history = []
if "current_result" not in st.session_state:
    st.session_state.current_result = None
if "current_input" not in st.session_state:
    st.session_state.current_input = ""
if "app_mode" not in st.session_state:
    st.session_state.app_mode = "review"  # 默认审校模式
if "gen_history" not in st.session_state:
    st.session_state.gen_history = []
    
# ================== 侧边栏：功能选择 ==================
st.sidebar.title("导航")
mode = st.sidebar.radio(
    "选择功能",
    ["📖 审校助手", "✍️ 内容生成", "🔄 版本对比", "📝 术语扫描"],
    index=0
)
if mode == "📖 审校助手":
    st.session_state.app_mode = "review"
elif mode == "✍️ 内容生成":
    st.session_state.app_mode = "generate"
elif mode == "🔄 版本对比":
    st.session_state.app_mode = "version_compare"
elif mode == "📝 术语扫描":
    st.session_state.app_mode = "term_scan"

# ================== 审校模式 ==================
def generate_review_word(result_dict):
    doc = Document()
    doc.add_heading('初中物理教材审校报告', 0)
    
    # 待审文本
    doc.add_heading('一、待审文本', level=1)
    doc.add_paragraph(result_dict.get('original_text', ''))
    
    # 知识审校
    if 'knowledge_review' in result_dict:
        doc.add_heading('二、知识审校（含版本对比）', level=1)
        doc.add_paragraph(result_dict['knowledge_review'])
    
    # 术语检查
    if 'terminology_review' in result_dict:
        doc.add_heading('三、术语检查', level=1)
        doc.add_paragraph(result_dict['terminology_review'])
    
    # 逻辑一致性
    if 'logic_review' in result_dict:
        doc.add_heading('四、逻辑一致性检查', level=1)
        doc.add_paragraph(result_dict['logic_review'])
    
    # 参考段落（可选）
    if 'retrieved_docs' in result_dict and result_dict['retrieved_docs']:
        doc.add_heading('五、参考段落（检索到的教材原文）', level=1)
        for i, doc_seg in enumerate(result_dict['retrieved_docs']):
            doc.add_heading(f'{i+1}. 【{doc_seg["version"]}】{doc_seg["chapter_name"]}', level=2)
            doc.add_paragraph(doc_seg['text'])
    
    # 保存到内存
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def render_review():
    # 侧边栏中的审校设置（仅在此模式下显示）
    with st.sidebar:
        st.markdown("---")
        st.header("⚙️ 审校设置")
        check_knowledge = st.checkbox("知识审校（含版本对比）", value=True)
        check_terminology = st.checkbox("术语检查", value=True)
        check_logic = st.checkbox("逻辑一致性检查", value=True)
        
        st.subheader("检索参数")
        top_k_per_version = st.number_input("每个版本取几个段落", min_value=1, max_value=5, value=2, step=1)
        max_results = st.number_input("最多返回段落数", min_value=5, max_value=30, value=15, step=5)
        
        st.markdown("---")
        st.header("📜 历史记录")
        if st.button("清空历史记录"):
            st.session_state.history = []
            st.session_state.current_result = None
            st.rerun()
        
        if st.session_state.history:
            for idx, item in enumerate(st.session_state.history):
                label = f"{item['timestamp']} - {item['input_preview'][:40]}..."
                if st.button(label, key=f"hist_{idx}"):
                    st.session_state.current_input = item["input_full"]
                    st.session_state.current_result = item["result"]
                    st.rerun()
        else:
            st.info("暂无历史记录")
    
    # 主区域
    st.title("📘 初中物理教材审校助手")
    
    # 输入方式
    tab1, tab2 = st.tabs(["✍️ 直接文本输入", "📁 上传文件"])
    input_text = ""
    with tab1:
        input_text = st.text_area("请输入待审校的教材文本", height=200, key="direct_input")
    with tab2:
        uploaded_file = st.file_uploader("上传 .txt 或 .md 文件", type=["txt", "md"])
        if uploaded_file is not None:
            try:
                input_text = uploaded_file.read().decode("utf-8")
            except UnicodeDecodeError:
                input_text = uploaded_file.read().decode("gbk", errors="ignore")
            st.text_area("文件内容预览", input_text, height=200, disabled=True)
    
    col1, col2 = st.columns([1, 5])
    with col1:
        run_btn = st.button("🚀 开始审校", type="primary", use_container_width=True)
    with col2:
        export_btn = st.button("📥 导出当前结果为 JSON", use_container_width=True)
    
    # 处理审校
    if run_btn and input_text.strip():
        with st.spinner("正在审校，请稍候...（可能需要 10-20 秒）"):
            try:
                result = review_text(
                    text=input_text,
                    check_knowledge=check_knowledge,
                    check_terminology=check_terminology,
                    check_logic=check_logic,
                    top_k_per_version=top_k_per_version,
                    max_results=max_results
                )
                st.session_state.current_result = result
                st.session_state.current_input = input_text
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                st.session_state.history.insert(0, {
                    "timestamp": timestamp,
                    "input_preview": input_text[:50],
                    "input_full": input_text,
                    "result": result
                })
                if len(st.session_state.history) > 20:
                    st.session_state.history = st.session_state.history[:20]
                st.rerun()
            except Exception as e:
                st.error(f"审校失败：{str(e)}")
    elif run_btn and not input_text.strip():
        st.warning("请输入或上传待审校的文本内容。")
    
    # 展示结果
    if st.session_state.current_result:
        result = st.session_state.current_result
        st.markdown("---")
        st.header("🔍 审校结果")
    
        # 创建标签页
        tab1, tab2, tab3, tab4 = st.tabs(["📖 知识审校", "📝 术语检查", "🔗 逻辑一致性", "📚 参考段落"])
    
        with tab1:
            if check_knowledge and "knowledge_review" in result:
                st.markdown(result["knowledge_review"])
            else:
                st.info("未开启知识审校或无结果。")
    
        with tab2:
            if check_terminology and "terminology_review" in result:
                st.markdown(result["terminology_review"])
            else:
                st.info("未开启术语检查或无结果。")
    
        with tab3:
            if check_logic and "logic_review" in result:
                st.markdown(result["logic_review"])
            else:
                st.info("未开启逻辑一致性检查或无结果。")
    
        with tab4:
            if "retrieved_docs" in result and result["retrieved_docs"]:
                for i, doc in enumerate(result["retrieved_docs"]):
                    st.markdown(f"**{i+1}. 【{doc['version']}】{doc['chapter_name']}**")
                    st.text(doc["text"][:800] + ("..." if len(doc["text"]) > 800 else ""))
                    st.markdown("---")
            else:
                st.info("无参考段落。")
    
        # 导出 Word 按钮（放在所有标签页之后）
        st.markdown("---")
        col_export, _ = st.columns([1, 5])
        with col_export:
            word_file = generate_review_word(result)
            st.download_button(
                label="📥 导出审校报告 (Word)",
                data=word_file,
                file_name=f"审校报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
        # 错误信息（如果有）
        if "error" in result:
            st.error(result["error"])

    # 导出 JSON
    if export_btn and st.session_state.current_result:
        export_data = {
            "timestamp": datetime.now().isoformat(),
            "input_text": st.session_state.current_input,
            "settings": {
                "check_knowledge": check_knowledge,
                "check_terminology": check_terminology,
                "check_logic": check_logic,
                "top_k_per_version": top_k_per_version,
                "max_results": max_results
            },
            "result": st.session_state.current_result
        }
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        st.download_button(
            label="📥 点击下载 JSON 文件",
            data=json_str,
            file_name=f"审校结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json"
        )
    elif export_btn:
        st.warning("没有可导出的结果，请先进行审校。")

# ================== 内容生成模式 ==================
def render_content_generation():
    st.title("✍️ 辅助内容生成")
    st.markdown("利用大模型辅助教材内容编写，支持习题批阅和习题生成。")
    
    # 创建三个标签页：习题批阅、习题生成、历史记录
    tab1, tab2, tab3 = st.tabs(["📝 习题批阅", "✨ 习题生成", "📜 历史记录"])
    
    # ========== 习题批阅 ==========
    with tab1:
        st.subheader("习题批阅")
        question = st.text_area("请输入题目", height=150, key="grading_question")
        student_answer = st.text_area("请输入学生答案", height=100, key="student_answer")
        if st.button("开始批阅", key="grade_btn"):
            if question and student_answer:
                with st.spinner("批阅中..."):
                    prompt = f"""请批阅以下初中物理习题，判断学生答案是否正确，并给出详细解析。
如果答案错误，请指出错误原因，并给出正确答案。
题目：{question}
学生答案：{student_answer}"""
                    result = call_deepseek(prompt)
                    st.markdown("### 批阅结果")
                    st.markdown(result)
                    # 保存历史
                    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    st.session_state.gen_history.insert(0, {
                        "type": "习题批阅",
                        "timestamp": timestamp,
                        "input_preview": question[:50] + "...",
                        "question": question,
                        "student_answer": student_answer,
                        "result": result
                    })
                    if len(st.session_state.gen_history) > 20:
                        st.session_state.gen_history = st.session_state.gen_history[:20]
            else:
                st.warning("请填写题目和学生答案")
    
    # ========== 习题生成 ==========
    with tab2:
        st.subheader("习题生成")
        col1, col2 = st.columns(2)
        with col1:
            topic = st.text_input("知识点", key="topic_gen", placeholder="例如：欧姆定律")
            question_type = st.selectbox("题型", ["选择题", "填空题", "计算题", "简答题"], key="q_type")
        with col2:
            difficulty = st.selectbox("难度", ["容易", "中等", "较难"], key="difficulty")
            num_questions = st.slider("题数", min_value=1, max_value=5, value=1, key="num_q")
        
        if st.button("生成习题", key="gen_btn"):
            if topic:
                with st.spinner("生成中..."):
                    prompt = f"""请生成{num_questions}道初中物理关于「{topic}」的{difficulty}难度{question_type}。
要求：
- 每道题单独列出，标注题号
- 包含题目和参考答案
- 题目应贴合初中生认知水平
- 输出格式清晰，便于阅读"""
                    result = call_deepseek(prompt)
                    st.markdown("### 生成的习题")
                    st.markdown(result)
                    # 保存历史
                    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    st.session_state.gen_history.insert(0, {
                        "type": "习题生成",
                        "timestamp": timestamp,
                        "input_preview": f"{topic} ({question_type}, {difficulty}, {num_questions}题)",
                        "topic": topic,
                        "question_type": question_type,
                        "difficulty": difficulty,
                        "num_questions": num_questions,
                        "result": result
                    })
                    if len(st.session_state.gen_history) > 20:
                        st.session_state.gen_history = st.session_state.gen_history[:20]
            else:
                st.warning("请输入知识点")
    
    # ========== 历史记录 ==========
    with tab3:
        st.subheader("内容生成历史")
        if st.button("清空历史记录"):
            st.session_state.gen_history = []
            st.rerun()
        if not st.session_state.gen_history:
            st.info("暂无历史记录")
        else:
            for idx, item in enumerate(st.session_state.gen_history):
                with st.expander(f"{item['timestamp']} - {item['type']}: {item['input_preview']}"):
                    if item['type'] == "习题批阅":
                        st.markdown(f"**题目**: {item['question']}")
                        st.markdown(f"**学生答案**: {item['student_answer']}")
                        st.markdown(f"**批阅结果**:\n{item['result']}")
                    else:
                        st.markdown(f"**生成参数**: {item['input_preview']}")
                        st.markdown(f"**生成结果**:\n{item['result']}")

# ================== 版本对比模式 ==================
def render_version_compare():
    st.title("🔄 跨版本智能对比器")
    st.markdown("选择两个教材版本的同一年级/册次/章节，系统将自动分析结构差异、知识点表述差异，并给出撰写建议。")
    
    # 加载数据
    data_file = Path("./enhanced_data.json")
    if not data_file.exists():
        st.error("enhanced_data.json 不存在")
        return
    with open(data_file, 'r', encoding='utf-8') as f:
        corpus = json.load(f)
    
    versions = sorted(set(item['version'] for item in corpus))
    grades = sorted(set(item['grade'] for item in corpus))
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("版本 A")
        ver1 = st.selectbox("版本", versions, key="ver1")
        grade1 = st.selectbox("年级", grades, key="grade1")
        volumes1 = sorted(set(item['volume'] for item in corpus if item['version'] == ver1 and item['grade'] == grade1))
        volume1 = st.selectbox("册次", volumes1, key="vol1")
        chapters1 = sorted(set(item['chapter_name'] for item in corpus if item['version'] == ver1 and item['grade'] == grade1 and item['volume'] == volume1))
        chap1 = st.selectbox("章节", chapters1, key="chap1")
    
    with col2:
        st.subheader("版本 B")
        ver2 = st.selectbox("版本", versions, key="ver2")
        grade2 = st.selectbox("年级", grades, key="grade2")
        volumes2 = sorted(set(item['volume'] for item in corpus if item['version'] == ver2 and item['grade'] == grade2))
        volume2 = st.selectbox("册次", volumes2, key="vol2")
        chapters2 = sorted(set(item['chapter_name'] for item in corpus if item['version'] == ver2 and item['grade'] == grade2 and item['volume'] == volume2))
        chap2 = st.selectbox("章节", chapters2, key="chap2")
    
    def get_content(ver, grade, volume, chap_name):
        for item in corpus:
            if item['version'] == ver and item['grade'] == grade and item['volume'] == volume and item['chapter_name'] == chap_name:
                return item['content']
        return None
    
    content1 = get_content(ver1, grade1, volume1, chap1)
    content2 = get_content(ver2, grade2, volume2, chap2)
    
    if not content1 or not content2:
        st.warning("未找到所选章节内容")
        return
    
    # 按钮：点击后才调用 API
    if st.button("开始智能对比", type="primary"):
        with st.spinner("正在分析两个版本的结构与表述差异，请稍候（约15-20秒）..."):
            prompt = f"""你是一位初中物理教材编辑专家。请对比以下两个版本的教材章节，从出版角度分析它们的异同，并给出撰写建议。

## 版本 A（{ver1} - {chap1}）
{content1}

## 版本 B（{ver2} - {chap2}）
{content2}

请按照以下格式输出：

### 一、结构对比
- 章节组织方式（子标题设置、内容展开顺序）的异同。
- 各自的特点（例如：版本A更强调实验引入，版本B更注重概念定义）。

### 二、知识点表述对比
- 选取2-3个核心知识点，对比它们的定义、公式、例子、图示等表述差异。
- 指出哪个版本的表述更严谨、哪个更生动，或各有优劣。

### 三、整体评价
- 从教学性、科学性、可读性三个维度简要评价两个版本。

### 四、撰写建议（针对该章节/专题）
- 结合两个版本的优点，给出一个理想的编写方案。
- 指出可以改进的地方（如增加某类例子、调整顺序、统一术语等）。

注意：输出使用Markdown格式，清晰分段，不要使用HTML标签。"""
            result = call_deepseek(prompt)
            st.session_state.compare_result = result  # 存储到 session_state
            st.markdown("### 对比结果")
            st.markdown(result)
    
    # 下载按钮：如果 session_state 中有结果，就显示
    if "compare_result" in st.session_state and st.session_state.compare_result:
        st.download_button(
            label="📥 导出对比报告 (Markdown)",
            data=st.session_state.compare_result,
            file_name=f"对比报告_{ver1}_{chap1}_vs_{ver2}_{chap2}.md",
            mime="text/markdown"
        )

# ================== 术语一致性扫描模式 ==================
def render_term_scan():
    st.title("📝 术语一致性扫描")
    st.markdown("上传教材稿件（.txt 或 .md），系统将自动检测其中使用的物理术语，并与标准术语词典对比，识别不规范术语和同义词混用。")
    
    input_text = ""
    tab1, tab2 = st.tabs(["✍️ 直接文本输入", "📁 上传文件"])
    with tab1:
        input_text = st.text_area("请输入待扫描的文本", height=200, key="term_input")
    with tab2:
        uploaded_file = st.file_uploader("上传 .txt 或 .md 文件", type=["txt", "md"])
        if uploaded_file is not None:
            try:
                input_text = uploaded_file.read().decode("utf-8")
            except UnicodeDecodeError:
                input_text = uploaded_file.read().decode("gbk", errors="ignore")
            st.text_area("文件内容预览", input_text, height=200, disabled=True)
    
    if st.button("开始扫描", type="primary"):
        if input_text.strip():
            with st.spinner("正在扫描术语，请稍候..."):
                from review_module import terminology_review
                result = terminology_review(input_text)
                st.markdown("### 扫描结果")
                st.markdown(result)
                # 提供导出选项
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="📥 导出结果为 Markdown",
                        data=result,
                        file_name="术语扫描结果.md",
                        mime="text/markdown"
                    )
                # 可选：导出 CSV 格式（将表格转换为 CSV）
                # 这里先不实现，因为结果中可能包含多个表格，简单起见只提供 Markdown
        else:
            st.warning("请输入或上传文本内容。")

# ================== 主逻辑 ==================
if st.session_state.app_mode == "review":
    render_review()
elif st.session_state.app_mode == "generate":
    render_content_generation()
elif st.session_state.app_mode == "version_compare":
    render_version_compare()
elif st.session_state.app_mode == "term_scan":
    render_term_scan()